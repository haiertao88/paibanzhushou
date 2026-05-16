
# 生成干净的 Streamlit 应用代码 - 不包含任何文件写入操作
# 直接输出到文件供下载

code = r'''#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WPS 规格书助手 - Web 全能版 V1.0
基于 Streamlit 的网页端完整移植
部署：streamlit run app.py
"""

import streamlit as st
import os, json, re, requests, tempfile, zipfile, io, base64, hashlib, random, time
from PIL import Image
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pdfplumber

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

# ═══════════════════════════════════════════
# 1. 页面配置与主题
# ═══════════════════════════════════════════
st.set_page_config(
    page_title="WPS 规格书助手 - Web全能版",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 深色科技主题 CSS
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(135deg, #1a1f2e 0%, #16213e 50%, #0f3460 100%);
    }
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1e293b 0%, #0f172a 100%);
        border-right: 1px solid #334155;
    }
    .main .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
        max-width: 100%;
    }
    .stButton > button {
        border-radius: 8px !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        border: none !important;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.3);
    }
    .stTextArea textarea {
        background: #0f172a !important;
        color: #e2e8f0 !important;
        border: 1px solid #334155 !important;
        border-radius: 8px !important;
        font-family: 'Consolas', monospace !important;
    }
    .stTextInput input {
        background: #0f172a !important;
        color: #e2e8f0 !important;
        border: 1px solid #334155 !important;
        border-radius: 6px !important;
    }
    .stSelectbox > div > div {
        background: #0f172a !important;
        color: #e2e8f0 !important;
        border: 1px solid #334155 !important;
    }
    .stTabs [data-baseweb="tab-list"] {
        background: rgba(30, 41, 59, 0.6) !important;
        border-radius: 8px !important;
        padding: 4px !important;
    }
    .stTabs [data-baseweb="tab"] {
        color: #94a3b8 !important;
        border-radius: 6px !important;
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
    }
    .stProgress > div > div {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%) !important;
    }
    hr {
        border-color: #334155 !important;
        margin: 12px 0 !important;
    }
    h1, h2, h3 {
        color: #f1f5f9 !important;
        font-weight: 700 !important;
    }
    h4, h5, h6 {
        color: #cbd5e1 !important;
        font-weight: 600 !important;
    }
    p, span, label {
        color: #e2e8f0 !important;
    }
    code {
        background: #1e293b !important;
        color: #38bdf8 !important;
        padding: 2px 6px !important;
        border-radius: 4px !important;
    }
    .status-dot {
        display: inline-block;
        width: 8px;
        height: 8px;
        border-radius: 50%;
        margin-right: 6px;
    }
    .status-ready { background: #38ef7d; box-shadow: 0 0 8px #38ef7d; }
    .status-busy { background: #fbbf24; box-shadow: 0 0 8px #fbbf24; }
    .status-error { background: #ef4444; box-shadow: 0 0 8px #ef4444; }
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }
    ::-webkit-scrollbar-track {
        background: #0f172a;
    }
    ::-webkit-scrollbar-thumb {
        background: #475569;
        border-radius: 4px;
    }
    ::-webkit-scrollbar-thumb:hover {
        background: #64748b;
    }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════
# 2. 全局常量与配置
# ═══════════════════════════════════════════
KIMI_API_URL = "https://api.moonshot.cn/v1/chat/completions"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
BAIDU_TRANS_URL = "https://fanyi-api.baidu.com/api/trans/vip/translate"

CN_HEADER_KEYWORDS = ["产品描述","产品特点","产品指标","应用场景","技术参数","产品介绍","安装方式","使用说明","注意事项","产品图片","产品包装","安装方法"]
EN_HEADER_KEYWORDS = ["Product Description","Product Features","Technical Specifications","Product Specifications","Applications","Application Scenarios","Instructions","Installation","Notes","Product Images","Product Packaging"]

FONT_CHOICES = ["微软雅黑","宋体","黑体","楷体","仿宋","Arial","Times New Roman","Calibri","Verdana","Georgia"]
BULLET_STYLES = {
    "● 实心圆": "● ", 
    "■ 实心方": "■ ", 
    "▶ 三角形": "▶ ", 
    "◆ 菱形": "◆ ", 
    "○ 空心圆": "○ ", 
    "① 带圈数字": "__NUM__", 
    "1. 数字编号": "__DOT__", 
    "— 短横线": "— "
}

THEME_COLORS = {
    "科技蓝": (41, 128, 185),
    "活力橙": (230, 126, 34),
    "中国红": (231, 76, 60),
    "环保绿": (39, 174, 96),
    "沉稳灰": (127, 140, 141)
}

URL_PATTERN = re.compile(r'https?://[^\s<>"\'\u3002\u3001\uff1b\uff1a\uff09\)\]\}]+')

# ═══════════════════════════════════════════
# 3. 状态管理
# ═══════════════════════════════════════════
def init_session_state():
    defaults = {
        'ai_provider': 'kimi',
        'kimi_key': '',
        'deepseek_key': '',
        'deepseek_model': 'deepseek-chat',
        'bd_id': '',
        'bd_key': '',
        'font_cn': "微软雅黑",
        'font_en': "Arial",
        'title_size': 14,
        'body_size': 11,
        'cover_cn': "产品规格书",
        'cover_en': "Product Specification",
        'bullet': "● 实心圆",
        'feature_brief': True,
        'bilingual_table': False,
        'line_spacing': 1.2,
        'theme_color': "科技蓝",
        'prod_name': 'XXXX',
        'custom_prompt': '',
        'use_web_crawler': False,
        'urls': '',
        'txt_cn': '',
        'txt_en': '',
        'raw_text': '',
        'gallery_images': [],
        'selected_images': set(),
        'bg_cover_bytes': None,
        'bg_body_bytes': None,
        'current_lang': 'cn',
        'ai_generating': False,
        'translating': False,
        'log_messages': [],
        'undo_stack': []
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()

# ═══════════════════════════════════════════
# 4. 核心功能函数
# ═══════════════════════════════════════════
def log(msg, type="info"):
    timestamp = time.strftime("%H:%M:%S")
    st.session_state.log_messages.append((timestamp, msg, type))
    if len(st.session_state.log_messages) > 100:
        st.session_state.log_messages = st.session_state.log_messages[-100:]

def clean_markdown(t):
    if not t: return ""
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)
    t = re.sub(r'^#+\s*', '', t)
    return t.strip()

def is_header(raw_line):
    clean = clean_markdown(raw_line)
    if clean in CN_HEADER_KEYWORDS or clean in EN_HEADER_KEYWORDS: return True
    if re.match(r'^\*\*.+\*\*$', raw_line) and not raw_line.startswith(('-','*','|')): return True
    return False

def get_bullet_symbol():
    sym = BULLET_STYLES.get(st.session_state.bullet, "● ")
    if sym in ("__NUM__", "__DOT__"): sym = "• "
    return sym

# ── AI 调用 ──
def call_ai_api(prompt, system_msg="你是资深的规格书编辑。"):
    provider = st.session_state.get('ai_provider', 'kimi')
    try:
        if provider == 'kimi':
            key = st.session_state.get('kimi_key', '')
            if not key: return "【错误】请配置 Kimi API Key"
            res = requests.post(KIMI_API_URL, 
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json={"model": "moonshot-v1-8k", "messages": [
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ], "temperature": 0.2},
                timeout=60)
        else:
            key = st.session_state.get('deepseek_key', '')
            model = st.session_state.get('deepseek_model', 'deepseek-chat')
            if not key: return "【错误】请配置 DeepSeek API Key"
            res = requests.post(DEEPSEEK_API_URL,
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json={"model": model, "messages": [
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ], "temperature": 0.2, "stream": False},
                timeout=60)
        
        data = res.json()
        if 'choices' in data and len(data['choices']) > 0:
            return data['choices'][0]['message']['content']
        return f"API 错误: {data.get('error', '未知错误')}"
    except Exception as e:
        return f"请求错误: {str(e)}"

# ── 百度翻译 ──
def call_baidu_translate(texts, src, tgt):
    if not texts: return []
    appid = st.session_state.get('bd_id', '')
    appkey = st.session_state.get('bd_key', '')
    if not appid or not appkey:
        return ["【错误：未配置百度翻译API】"] * len(texts)
    
    results = []
    MAX_BYTES = 4000
    batches = []
    current_batch = []
    current_len = 0
    
    for t in texts:
        clean_t = str(t).replace('\n', ' ').replace('\r', ' ')
        t_len = len(clean_t.encode('utf-8'))
        if current_len + t_len > MAX_BYTES and current_batch:
            batches.append(current_batch)
            current_batch = [clean_t]
            current_len = t_len
        else:
            current_batch.append(clean_t)
            current_len += t_len + 1
    if current_batch:
        batches.append(current_batch)
    
    for batch in batches:
        query = '\n'.join(batch)
        salt = str(random.randint(32768, 65536))
        sign_str = appid + query + salt + appkey
        sign = hashlib.md5(sign_str.encode("utf-8")).hexdigest()
        
        try:
            resp = requests.post(BAIDU_TRANS_URL, data={
                "q": query, "from": src, "to": tgt,
                "appid": appid, "salt": salt, "sign": sign
            }, timeout=30)
            data = resp.json()
            if "error_code" in data and data["error_code"] != "52000":
                return [f"翻译错误[{data['error_code']}]"] * len(texts)
            for item in data.get("trans_result", []):
                results.append(item.get("dst", ""))
        except Exception as e:
            return [f"翻译请求错误: {e}"] * len(texts)
    
    if len(results) < len(texts):
        results.extend([""] * (len(texts) - len(results)))
    return results[:len(texts)]

# ── 网页爬虫 ──
class WebFetcher:
    UA = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    
    @staticmethod
    def fetch(url, max_chars=8000):
        headers = {'User-Agent': WebFetcher.UA, 'Accept': 'text/html,*/*;q=0.8'}
        try:
            resp = requests.get(url, headers=headers, timeout=30, verify=False)
            resp.encoding = resp.apparent_encoding or 'utf-8'
            html = resp.text
            if len(html) < 100: return "", "", "内容过短"
        except Exception as e:
            return "", "", f"请求失败: {str(e)[:80]}"
        
        body, table = "", ""
        if HAS_BS4:
            try:
                soup = BeautifulSoup(html, 'html.parser')
                for tag in soup(['script','style','nav','footer','header','aside','iframe','noscript','form']):
                    tag.decompose()
                main = soup.find('main') or soup.find('article') or soup.find(class_=re.compile(r'product|detail|content',re.I)) or soup.body or soup
                body = main.get_text(separator='\n', strip=True)
                
                tparts = []
                for tbl in soup.find_all('table'):
                    md = []
                    for row in tbl.find_all('tr'):
                        cells = [c.get_text(strip=True).replace('|','/').replace('\n',' ') for c in row.find_all(['td','th'])]
                        if any(cells): md.append("| "+" | ".join(cells)+" |")
                    if len(md) >= 2:
                        nc = md[0].count('|')-1
                        tparts.append(md[0]+"\n|"+" | ".join(["---"]*max(nc,1))+" |\n"+"\n".join(md[1:]))
                if tparts: table = "\n[规格表格]\n"+"\n\n".join(tparts)
            except:
                body = re.sub(r'<[^>]+>','\n',re.sub(r'<script.*?</script>','',html,flags=re.DOTALL|re.I))
        else:
            body = re.sub(r'<script.*?</script>','',html,flags=re.DOTALL|re.I)
            body = re.sub(r'<style.*?</style>','',body,flags=re.DOTALL|re.I)
            body = re.sub(r'<[^>]+>','\n',body)
        
        lines = [l.strip() for l in body.split('\n') if l.strip() and len(l.strip())>2]
        body = '\n'.join(lines)[:max_chars]
        table = table[:3000]
        return body, table, ""
    
    @staticmethod
    def fetch_multi(urls, max_total=10000):
        ab, at, ae, af = [], [], [], []
        tl = 0
        for url in urls[:5]:
            b, t, e = WebFetcher.fetch(url, max_chars=max(3000,(max_total-tl)//max(1,5-len(af))))
            if e: ae.append(f"[{url}] {e}")
            if b or t:
                ab.append(f"\n来源:{url}\n{b}"); at.append(t); af.append(url)
                tl += len(b)+len(t)
            if tl >= max_total: break
        return '\n'.join(ab), '\n'.join(at), ae, af

# ── 图片提取 ──
def extract_images_from_file(uploaded_file):
    imgs = []
    try:
        if uploaded_file.name.endswith(".pdf"):
            import fitz
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            for i in range(len(doc)):
                for img in doc[i].get_images(full=True):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    if len(image_bytes) > 3000: imgs.append((image_bytes, f"pdf_img_{i}_{xref}.png"))
            doc.close()
        elif uploaded_file.name.endswith(".docx"):
            with zipfile.ZipFile(io.BytesIO(uploaded_file.read())) as z:
                for f in z.namelist():
                    if f.startswith("word/media/") and f.lower().endswith(('.png','.jpg','.jpeg','.webp')):
                        data = z.read(f)
                        if len(data) > 3000: imgs.append((data, os.path.basename(f)))
    except Exception as e:
        log(f"图片提取错误: {e}", "error")
    return imgs

# ── 构建 AI Prompt ──
def build_ai_prompt(lang):
    if lang == 'cn':
        bilingual_rule = ""
        if st.session_state.get('bilingual_table', False):
            bilingual_rule = "【特殊要求：产品指标的 Markdown 表格，对于每一个参数项（包括表头），必须固定输出两行，上一行保持英文原文，紧接着的下一行输出对应的中文翻译对照。】"
        
        rules = f"""【文案生成严格要求】
1. 剔除资料中所有的品牌名称。
2. 必须按顺序包含以下章节标题（所有标题必须加粗，且绝对不能带有数字序号）：
   **产品描述** / **产品特点** / **产品指标** / **应用场景** / **安装方式** / **使用说明**
3. 产品描述：必须在第一句话明确点出【具体产品名称及适用线型/接口】。重点描述核心用途和基本功能。【绝对不要提及具体材质名称】，只需写"采用高品质金属材料和精密加工"。字数80-120字。
4. 产品特点：不少于5个。每条严格采用"• 关键词：简短说明"格式，简短说明控制在10个字以内，不要包含任何具体参数数值。
5. 产品指标：必须使用 Markdown 表格。务必100%完整提取所有参数指标，包括电气性能和机械/材质属性。{bilingual_rule}
6. 应用场景：不少于4个。每条严格采用"• "开头。
7. 使用说明：放在安装方式之后。提供3-5点常规指导，采用"• "开头，通俗易懂，不要涉及任何具体参数数值。
8. 所有内容必须使用中文。"""
    else:
        rules = """【STRICT COPYWRITING RULES】
1. Remove all brand names.
2. MUST sequentially include: **Product Description** / **Product Features** / **Product Specifications** / **Application Scenarios** / **Installation** / **Instructions for Use** (bolded, no numbers)
3. Product Description: MUST state specific [Product Name and Cable Type/Interface] in first sentence. Focus on core purpose. DO NOT mention specific materials. Use "high-quality metal materials and precision machining". 80-120 words.
4. Product Features: At least 5. Format: "• Keyword: Short description" (under 10 words, NO parameters).
5. Product Specifications: MUST be Markdown table. Extract 100% of all parameters including Electrical and Material/Mechanical properties.
6. Application Scenarios: At least 4. Use "• " prefix.
7. Instructions for Use: After Installation. 3-5 bullet points, simple and accessible. NO parameters.
8. ALL content MUST be in English."""
    
    parts = [rules]
    
    # 网页爬虫
    if st.session_state.get('use_web_crawler', False):
        raw_url = st.session_state.get('urls', '')
        urls = [u for u in URL_PATTERN.findall(raw_url) if len(u)>15]
        if urls:
            log(f"正在爬取 {len(urls)} 个网页...", "info")
            b, t, e, f = WebFetcher.fetch_multi(urls)
            if b: parts.append(f"\n【网页正文资料】\n{b}")
            if t: parts.append(f"\n【网页表格资料】\n{t}")
            if f: log(f"网页提取成功: {len(f)} 个来源", "success")
    
    # 自定义指令
    custom = st.session_state.get('custom_prompt', '').strip()
    if custom and custom != "请输入自定义指令...":
        parts.append(f"\n【用户附加指令】\n{custom}")
    
    # 本地文档
    if st.session_state.raw_text:
        parts.append(f"\n【本地文档资料】\n{st.session_state.raw_text[:6000]}")
    
    if not st.session_state.raw_text and not st.session_state.get('use_web_crawler', False):
        parts.append("\n【请基于通用通信产品知识生成规格书】")
    
    return '\n'.join(parts)

# ── Word 生成 ──
def _add_word_bg(section, bg_bytes):
    if not bg_bytes: return
    hdr = section.header
    hdr.is_linked_to_previous = False
    if len(hdr.paragraphs) == 0:
        p = hdr.add_paragraph()
    else:
        p = hdr.paragraphs[0]
        p.text = ""
        for r in p.runs: r._element.getparent().remove(r._element)
    
    run = p.add_run()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        tmp.write(bg_bytes)
        tmp_path = tmp.name
    
    try:
        pic = run.add_picture(tmp_path, width=Mm(210), height=Mm(297))
        anc = OxmlElement('wp:anchor')
        anc.set(qn('wp:behindDoc'), '1')
        for tag, val in [('wp:positionH', 'page'), ('wp:positionV', 'page')]:
            pos = OxmlElement(tag)
            pos.set(qn('wp:relativeFrom'), val)
            off = OxmlElement('wp:posOffset'); off.text = '0'
            pos.append(off); anc.append(pos)
        for child in pic._inline: anc.append(child)
        pic._inline.getparent().replace(pic._inline, anc)
    finally:
        os.unlink(tmp_path)

def generate_word_document(lang):
    doc = Document()
    fn = st.session_state.font_cn if lang == 'cn' else st.session_state.font_en
    txt = st.session_state.txt_cn if lang == 'cn' else st.session_state.txt_en
    
    # 页面设置
    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Mm(25.4); sec.bottom_margin = Mm(25.4); sec.left_margin = Mm(19.1); sec.right_margin = Mm(19.1)
    doc.styles['Normal'].font.name = fn
    if lang == 'cn':
        doc.styles['Normal']._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), fn)
    
    # 封面
    if st.session_state.bg_cover_bytes:
        _add_word_bg(sec, st.session_state.bg_cover_bytes)
    
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(10): cp.add_run('\n')
    cover_title = st.session_state.cover_cn if lang == 'cn' else st.session_state.cover_en
    rn = cp.add_run(cover_title if cover_title else ("Product Specification" if lang=='en' else "产品规格书"))
    rn.bold = True; rn.font.size = Pt(max(st.session_state.title_size + 14, 24)); rn.font.name = fn
    if lang == 'cn': rn._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), fn)
    
    # 产品名称注入
    prod = st.session_state.get('prod_name', '')
    if prod and prod != 'XXXX':
        cp2 = doc.add_paragraph()
        cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn2 = cp2.add_run(prod)
        rn2.font.size = Pt(20); rn2.font.name = fn; rn2.bold = True
        if lang == 'cn': rn2._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), fn)
    
    # 正文节
    bs = doc.add_section()
    bs.header.is_linked_to_previous = False
    if st.session_state.bg_body_bytes:
        _add_word_bg(bs, st.session_state.bg_body_bytes)
    else:
        if len(bs.header.paragraphs) > 0:
            bs.header.paragraphs[0].text = ""
    
    # 行距设置
    line_spacing = st.session_state.get('line_spacing', 1.2)
    
    lines = txt.split('\n')
    tr = []
    
    def render_run(pa, text, size=st.session_state.body_size, bold=False, color=None):
        r = pa.add_run(text)
        r.font.size = Pt(size); r.font.name = fn; r.bold = bold
        if color: r.font.color.rgb = RGBColor(*color)
        if lang == 'cn': r._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), fn)
        return r
    
    for idx, line in enumerate(lines):
        lt = line.strip()
        if not lt: continue
        
        # 表格
        if '|' in lt:
            if '---' in lt: continue
            cells = [clean_markdown(c) for c in lt.strip('|').split('|') if c.strip() != '']
            if cells: tr.append(cells)
            nxt = lines[idx+1].strip() if idx+1 < len(lines) else ""
            if not nxt or '|' not in nxt:
                if tr:
                    t = doc.add_table(rows=len(tr), cols=max(len(r) for r in tr))
                    t.style = 'Table Grid'
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    # 主题色
                    theme = st.session_state.get('theme_color', '科技蓝')
                    r_val, g_val, b_val = THEME_COLORS.get(theme, (41, 128, 185))
                    
                    for r_idx, row_data in enumerate(tr):
                        for c_idx, cell_val in enumerate(row_data):
                            if c_idx < len(t.columns):
                                cell = t.cell(r_idx, c_idx)
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.paragraph_format.line_spacing = line_spacing
                                is_header_row = (r_idx == 0)
                                render_run(p, cell_val, st.session_state.body_size, is_header_row)
                                if is_header_row:
                                    shading = cell._tc.get_or_add_tcPr()
                                    shade = OxmlElement('w:shd')
                                    shade.set(qn('w:fill'), f'{r_val:02x}{g_val:02x}{b_val:02x}')
                                    shading.append(shade)
                                    for run in p.runs:
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                    tr = []; doc.add_paragraph()
            continue
        
        # 标题
        if is_header(lt):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = line_spacing
            render_run(p, clean_markdown(lt), st.session_state.title_size, True)
            continue
        
        # 图片占位
        if 'IMG_FRAME' in lt:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            render_run(p, "【产品图片排版预留区】", st.session_state.body_size, True, (100, 100, 100))
            continue
        
        # 普通文本
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = line_spacing
        if lt.startswith(('-', '*', '•')):
            sym = get_bullet_symbol()
            render_run(p, f"{sym}{clean_markdown(lt[1:])}")
        else:
            render_run(p, clean_markdown(lt))
    
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# ═══════════════════════════════════════════
# 5. 微调工具函数
# ═══════════════════════════════════════════
def remove_descriptions(text):
    """删除冒号后的说明文字"""
    lines = text.split('\n')
    for i, l in enumerate(lines):
        m = re.search(r'^(.*?)(：|:)', l)
        if m:
            prefix = m.group(1)
            clean_prefix = re.sub(r'^[-*•●■▶◆○①-⑳0-9.、\s]+', '', prefix)
            if len(clean_prefix) <= 30:
                lines[i] = l[:m.end()-1]
    return '\n'.join(lines)

def replace_numbers(text):
    """将序号替换为项目符号"""
    lines = text.split('\n')
    for i, l in enumerate(lines):
        new_l = re.sub(r'^(\d+[\.、]|[①-⑳])\s*', '• ', l.lstrip())
        if new_l != l.lstrip():
            indent = l[:len(l)-len(l.lstrip())]
            lines[i] = indent + new_l
    return '\n'.join(lines)

def remove_blank_lines(text):
    """删除空行"""
    lines = [l for l in text.split('\n') if l.strip()]
    return '\n'.join(lines)

def bold_prefix(text):
    """加粗冒号前的关键词"""
    lines = text.split('\n')
    result = []
    for l in lines:
        pos = l.find('：') if '：' in l else l.find(':')
        if pos != -1 and pos > 0:
            result.append(f"**{l[:pos+1]}**{l[pos+1:]}")
        else:
            result.append(l)
    return '\n'.join(result)

def ai_rewrite(text, style, lang='cn'):
    """AI 重写/精炼/详细/续行"""
    if not text.strip():
        return text
    
    en_chars = len(re.findall(r'[a-zA-Z]', text))
    cn_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    is_en = en_chars > cn_chars
    
    if style == "详细":
        if is_en:
            prompt = f"Please moderately expand the following excerpt. Add reasonable technical details. Control word count. Output directly without explanation:\n\n{text}"
        else:
            prompt = f"对以下片段进行适度扩写，增加合理的技术修饰词或细节，直接输出结果不解释：\n\n{text}"
    elif style == "精炼":
        if is_en:
            prompt = f"Please refine and condense the following excerpt. Remove redundant words. Output directly:\n\n{text}"
        else:
            prompt = f"对以下片段进行精简，剔除冗余词汇，直接输出：\n\n{text}"
    elif style == "续行":
        lines = [l for l in text.split('\n') if l.strip()]
        context = '\n'.join(lines[-2:]) if len(lines) >= 2 else text
        if is_en:
            prompt = f"Continue the following product features. Use the same bullet style. ONLY output the next line:\n\n{context}"
        else:
            prompt = f"续写以下产品特征，保持相同的项目符号风格，只输出紧接着的下一行：\n\n{context}"
    else:  # 重写
        if is_en:
            prompt = f"Please polish and rewrite the following excerpt professionally:\n\n{text}"
        else:
            prompt = f"请重新润色和重写以下片段，使其专业流畅：\n\n{text}"
    
    return call_ai_api(prompt, "你是资深的通信规格书编辑。")

# ═══════════════════════════════════════════
# 6. 侧边栏 - 系统设置
# ═══════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ 系统设置")
    
    with st.expander("🔑 API 密钥配置", expanded=False):
        st.selectbox("AI 提供商", ["kimi", "deepseek"], key="ai_provider")
        st.text_input("Kimi API Key", key="kimi_key", type="password")
        st.text_input("DeepSeek API Key", key="deepseek_key", type="password")
        st.text_input("DeepSeek 模型", key="deepseek_model", value="deepseek-chat")
        st.divider()
        st.text_input("百度翻译 APP ID", key="bd_id")
        st.text_input("百度翻译 密钥", key="bd_key", type="password")
    
    with st.expander("🎨 排版参数", expanded=False):
        st.selectbox("中文字体", FONT_CHOICES, key="font_cn")
        st.selectbox("英文字体", FONT_CHOICES, key="font_en")
        c1, c2 = st.columns(2)
        c1.number_input("标题字号", 8, 36, key="title_size")
        c2.number_input("正文字号", 6, 24, key="body_size")
        st.selectbox("项目符号", list(BULLET_STYLES.keys()), key="bullet")
        st.selectbox("主题色系", list(THEME_COLORS.keys()), key="theme_color")
        st.slider("行距倍数", 1.0, 2.0, 1.2, 0.05, key="line_spacing")
        st.checkbox("产品特点含简短说明", key="feature_brief")
        st.checkbox("表格中英双语对照", key="bilingual_table")
    
    st.divider()
    st.markdown("### 📊 系统状态")
    if st.session_state.ai_generating:
        st.markdown('<span class="status-dot status-busy"></span>AI 生成中...', unsafe_allow_html=True)
    elif st.session_state.translating:
        st.markdown('<span class="status-dot status-busy"></span>翻译中...', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-dot status-ready"></span>系统就绪', unsafe_allow_html=True)
    
    # 日志显示
    st.divider()
    st.markdown("### 📝 操作日志")
    log_container = st.container(height=200)
    with log_container:
        for ts, msg, typ in st.session_state.log_messages[-20:]:
            color = "#38ef7d" if typ == "success" else "#fbbf24" if typ == "warning" else "#ef4444" if typ == "error" else "#94a3b8"
            st.markdown(f"<small style='color:{color}'>[{ts}] {msg}</small>", unsafe_allow_html=True)

# ═══════════════════════════════════════════
# 7. 主界面 - 三栏布局
# ═══════════════════════════════════════════
st.markdown("<h1 style='text-align: center; margin-bottom: 20px;'>📄 WPS 规格书助手 - Web 全能版</h1>", unsafe_allow_html=True)

# 顶部快捷工具栏
toolbar_col1, toolbar_col2, toolbar_col3, toolbar_col4, toolbar_col5, toolbar_col6, toolbar_col7 = st.columns([1,1,1,1,1,1,1])

with toolbar_col1:
    if st.button("✂️ 去说明", use_container_width=True, help="删除冒号后的说明文字"):
        if st.session_state.current_lang == 'cn' and st.session_state.txt_cn:
            st.session_state.txt_cn = remove_descriptions(st.session_state.txt_cn)
            st.session_state.undo_stack.append(('cn', st.session_state.txt_cn))
            log("已删除说明文字", "success"); st.rerun()
        elif st.session_state.current_lang == 'en' and st.session_state.txt_en:
            st.session_state.txt_en = remove_descriptions(st.session_state.txt_en)
            st.session_state.undo_stack.append(('en', st.session_state.txt_en))
            log("已删除说明文字", "success"); st.rerun()

with toolbar_col2:
    if st.button("🔄 去序号", use_container_width=True, help="将数字序号替换为项目符号"):
        if st.session_state.current_lang == 'cn' and st.session_state.txt_cn:
            st.session_state.txt_cn = replace_numbers(st.session_state.txt_cn)
            log("已替换序号", "success"); st.rerun()
        elif st.session_state.current_lang == 'en' and st.session_state.txt_en:
            st.session_state.txt_en = replace_numbers(st.session_state.txt_en)
            log("已替换序号", "success"); st.rerun()

with toolbar_col3:
    if st.button("🗑️ 去空行", use_container_width=True, help="删除所有空行"):
        if st.session_state.current_lang == 'cn' and st.session_state.txt_cn:
            st.session_state.txt_cn = remove_blank_lines(st.session_state.txt_cn)
            log("已删除空行", "success"); st.rerun()
        elif st.session_state.current_lang == 'en' and st.session_state.txt_en:
            st.session_state.txt_en = remove_blank_lines(st.session_state.txt_en)
            log("已删除空行", "success"); st.rerun()

with toolbar_col4:
    if st.button("🅱️ 前加粗", use_container_width=True, help="加粗冒号前的关键词"):
        if st.session_state.current_lang == 'cn' and st.session_state.txt_cn:
            st.session_state.txt_cn = bold_prefix(st.session_state.txt_cn)
            log("已加粗关键词", "success"); st.rerun()
        elif st.session_state.current_lang == 'en' and st.session_state.txt_en:
            st.session_state.txt_en = bold_prefix(st.session_state.txt_en)
            log("已加粗关键词", "success"); st.rerun()

with toolbar_col5:
    rewrite_style = st.selectbox("", ["重写", "精炼", "详细", "续行"], label_visibility="collapsed", key="rewrite_style")

with toolbar_col6:
    if st.button("✨ AI改写", use_container_width=True, help="使用AI重写选中风格"):
        current_text = st.session_state.txt_cn if st.session_state.current_lang == 'cn' else st.session_state.txt_en
        if current_text:
            st.session_state.ai_generating = True
            result = ai_rewrite(current_text, rewrite_style, st.session_state.current_lang)
            st.session_state.ai_generating = False
            if st.session_state.current_lang == 'cn':
                st.session_state.txt_cn = result
            else:
                st.session_state.txt_en = result
            log(f"AI {rewrite_style} 完成", "success")
            st.rerun()

with toolbar_col7:
    if st.button("↩ 撤销", use_container_width=True, help="撤销上次操作"):
        if st.session_state.undo_stack:
            lang, text = st.session_state.undo_stack.pop()
            if lang == 'cn': st.session_state.txt_cn = text
            else: st.session_state.txt_en = text
            log("已撤销", "success")
            st.rerun()

st.divider()

# 三栏主布局
col_left, col_mid, col_right = st.columns([1.1, 2.2, 0.9], gap="medium")

# ═══════════════════════════════════════════
# 8. 左侧面板 - 资料与AI生成
# ═══════════════════════════════════════════
with col_left:
    st.markdown("### 🛠️ 核心操作")
    
    # 资料导入区
    with st.container(border=True):
        st.markdown("**📂 资料导入**")
        doc_file = st.file_uploader("上传 PDF/DOCX", type=['pdf', 'docx'], label_visibility="collapsed", key="doc_uploader")
        
        c1, c2 = st.columns(2)
        if c1.button("🚀 提取资料", use_container_width=True):
            if doc_file:
                with st.spinner("深度解析中..."):
                    # 提取文本
                    if doc_file.name.endswith(".pdf"):
                        try:
                            with pdfplumber.open(doc_file) as pdf:
                                texts = []
                                for p in pdf.pages:
                                    t = p.extract_text()
                                    if t: texts.append(t)
                                st.session_state.raw_text = "\n".join(texts)
                                log(f"PDF 文本提取完成: {len(st.session_state.raw_text)} 字", "success")
                        except Exception as e:
                            log(f"PDF 提取失败: {e}", "error")
                    else:
                        try:
                            docx_doc = Document(doc_file)
                            texts = [p.text for p in docx_doc.paragraphs if p.text.strip()]
                            st.session_state.raw_text = "\n".join(texts)
                            log(f"DOCX 文本提取完成: {len(st.session_state.raw_text)} 字", "success")
                        except Exception as e:
                            log(f"DOCX 提取失败: {e}", "error")
                    
                    # 提取图片
                    doc_file.seek(0)
                    extracted = extract_images_from_file(doc_file)
                    for img_bytes, img_name in extracted:
                        st.session_state.gallery_images.append((img_bytes, img_name))
                    log(f"提取 {len(extracted)} 张图片", "success")
            else:
                st.warning("请先上传文档")
        
        if c2.button("🗑 清空资料", use_container_width=True):
            st.session_state.raw_text = ""
            log("资料已清空", "warning")
            st.rerun()
        
        if st.session_state.raw_text:
            st.success(f"已加载: {len(st.session_state.raw_text)} 字")
    
    # 网页爬虫
    with st.container(border=True):
        st.markdown("**🌐 网页爬虫**")
        st.checkbox("启用网页爬虫", key="use_web_crawler")
        st.text_area("网址 (每行一个)", key="urls", height=60, placeholder="https://example.com/product")
    
    # 产品名称
    with st.container(border=True):
        st.markdown("**🏷️ 产品名称**")
        st.text_input("", key="prod_name", placeholder="输入产品名称")
    
    # AI 生成
    with st.container(border=True):
        st.markdown("**🤖 AI 文案生成**")
        
        # 预设指令
        preset = st.selectbox("预设模板", [
            "预设1：标准规格书",
            "预设2：侧重应用场景", 
            "预设3：侧重营销包装",
            "自定义指令"
        ], key="preset_template")
        
        if preset == "预设1：标准规格书":
            default_prompt = "生成标准通信产品规格书，包含产品描述、特点、指标表格、应用场景、安装方式、使用说明。"
        elif preset == "预设2：侧重应用场景":
            default_prompt = "着重梳理并补充产品的应用场景、适用范围和核心优势。"
        elif preset == "预设3：侧重营销包装":
            default_prompt = "使用更具吸引力和营销感的口吻对文案进行包装。"
        else:
            default_prompt = st.session_state.get('custom_prompt', '')
        
        st.text_area("自定义指令", value=default_prompt, key="custom_prompt", height=80)
        
        c1, c2 = st.columns(2)
        if c1.button("⚡ 生成中文", type="primary", use_container_width=True):
            st.session_state.ai_generating = True
            log("开始生成中文文案...", "info")
            prompt = build_ai_prompt('cn')
            result = call_ai_api(prompt)
            st.session_state.ai_generating = False
            if not result.startswith("【错误"):
                st.session_state.txt_cn = result
                st.session_state.current_lang = 'cn'
                # 自动提取产品名称
                try:
                    name_prompt = f"请从产品文案中提取产品名称（仅输出名称，10字以内）：\n{result[:500]}"
                    name = call_ai_api(name_prompt, "你是产品经理。").strip().strip('"').strip('"').strip('"')
                    if name and len(name) <= 15:
                        st.session_state.prod_name = name
                except: pass
                log("中文文案生成完成", "success")
            else:
                log(result, "error")
            st.rerun()
        
        if c2.button("⚡ 生成英文", type="primary", use_container_width=True):
            st.session_state.ai_generating = True
            log("开始生成英文文案...", "info")
            prompt = build_ai_prompt('en')
            result = call_ai_api(prompt)
            st.session_state.ai_generating = False
            if not result.startswith("【错误"):
                st.session_state.txt_en = result
                st.session_state.current_lang = 'en'
                log("英文文案生成完成", "success")
            else:
                log(result, "error")
            st.rerun()
    
    # 翻译
    with st.container(border=True):
        st.markdown("**🌐 极速翻译**")
        trans_dir = st.selectbox("翻译方向", ["英→中", "中→英", "自动→中", "自动→英"], key="trans_dir")
        
        if st.button("🚀 翻译当前文案", use_container_width=True):
            current_text = st.session_state.txt_cn if st.session_state.current_lang == 'cn' else st.session_state.txt_en
            if not current_text:
                st.warning("当前没有文案内容")
            else:
                st.session_state.translating = True
                log("开始翻译...", "info")
                
                # 解析方向
                dir_map = {"英→中": ("en", "zh"), "中→英": ("zh", "en"), "自动→中": ("auto", "zh"), "自动→英": ("auto", "en")}
                src, tgt = dir_map.get(trans_dir, ("en", "zh"))
                
                # 分段翻译
                lines = current_text.split('\n')
                texts = [l for l in lines if l.strip()]
                prefixes = []
                clean_texts = []
                for l in texts:
                    m = re.match(r'^([-•●■▶◆○—]\s*|\d+[\.、]\s*|[①-⑳]\s*)(.*)', l)
                    if m:
                        prefixes.append(m.group(1))
                        clean_texts.append(m.group(2))
                    else:
                        prefixes.append("")
                        clean_texts.append(l)
                
                translations = call_baidu_translate(clean_texts, src, tgt)
                result_lines = [p + t for p, t in zip(prefixes, translations)]
                result = '\n'.join(result_lines)
                
                st.session_state.translating = False
                if st.session_state.current_lang == 'cn':
                    st.session_state.txt_en = result
                    log("中文→英文翻译完成", "success")
                else:
                    st.session_state.txt_cn = result
                    log("英文→中文翻译完成", "success")
                st.rerun()

# ═══════════════════════════════════════════
# 9. 中间面板 - 编辑器与预览
# ═══════════════════════════════════════════
with col_mid:
    # 标签页切换
    tab_edit, tab_preview = st.tabs(["✏️ 文案编辑", "👁️ 实时预览"])
    
    with tab_edit:
        st.markdown("### 📝 Markdown 编辑器")
        
        # 语言切换
        lang_cols = st.columns([1,1,2])
        if lang_cols[0].button("🇨🇳 中文", use_container_width=True, type="secondary" if st.session_state.current_lang=='en' else "primary"):
            st.session_state.current_lang = 'cn'
            st.rerun()
        if lang_cols[1].button("🇺🇸 English", use_container_width=True, type="secondary" if st.session_state.current_lang=='cn' else "primary"):
            st.session_state.current_lang = 'en'
            st.rerun()
        
        # 编辑区
        if st.session_state.current_lang == 'cn':
            st.text_area("中文内容", key="txt_cn", height=500, label_visibility="collapsed")
        else:
            st.text_area("English Content", key="txt_en", height=500, label_visibility="collapsed")
        
        # 底部操作
        op1, op2, op3, op4 = st.columns([1,1,1,1])
        if op1.button("📸 +1图框"):
            if st.session_state.current_lang == 'cn':
                st.session_state.txt_cn += "\n\n[IMG_FRAME:1]"
            else:
                st.session_state.txt_en += "\n\n[IMG_FRAME:1]"
            st.rerun()
        if op2.button("📸 +2图框"):
            if st.session_state.current_lang == 'cn':
                st.session_state.txt_cn += "\n\n[IMG_FRAME:2]"
            else:
                st.session_state.txt_en += "\n\n[IMG_FRAME:2]"
            st.rerun()
        if op3.button("📥 导出Word"):
            st.info("请使用右侧导出按钮")
        if op4.button("📋 复制"):
            text = st.session_state.txt_cn if st.session_state.current_lang=='cn' else st.session_state.txt_en
            st.code(text, language="markdown")
    
    with tab_preview:
        st.markdown("### 👁️ A4 实时预览")
        
        preview_txt = st.session_state.txt_cn if st.session_state.current_lang == 'cn' else st.session_state.txt_en
        cover_title = st.session_state.cover_cn if st.session_state.current_lang == 'cn' else st.session_state.cover_en
        font_family = st.session_state.font_cn if st.session_state.current_lang == 'cn' else st.session_state.font_en
        title_px = st.session_state.title_size * 1.5
        body_px = st.session_state.body_size * 1.3
        
        # 构建 HTML 预览
        html_body = ""
        in_table = False
        
        for line in preview_txt.split('\n'):
            line = line.strip()
            if not line: continue
            
            if '|' in line:
                if '---' in line: continue
                if not in_table:
                    html_body += '<table style="width:100%; border-collapse:collapse; margin:10px 0; font-size:14px;">'
                    in_table = True
                cells = [clean_markdown(c) for c in line.strip('|').split('|') if c.strip() != '']
                html_body += "<tr>"
                for c in cells:
                    html_body += f'<td style="border:1px solid #333; padding:8px; text-align:center;">{c}</td>'
                html_body += "</tr>"
                continue
            
            if in_table:
                html_body += '</table>'
                in_table = False
            
            if 'IMG_FRAME' in line:
                html_body += '<div style="border:2px dashed #007AFF; padding:40px; text-align:center; margin:15px 0; background:rgba(0,122,255,0.05); color:#007AFF; border-radius:8px;"><b>🖼️ 产品图片排版预留区</b></div>'
                continue
            
            if is_header(line):
                html_body += f'<div style="margin-top:20px; margin-bottom:8px; font-size:{title_px}px; font-weight:bold; color:#1e293b;">{clean_markdown(line)}</div>'
                continue
            
            if line.startswith(('-', '*', '•')):
                sym = get_bullet_symbol()
                html_body += f'<div style="margin-bottom:6px; margin-left:20px; font-size:{body_px}px;">{sym} {clean_markdown(line[1:])}</div>'
            else:
                html_body += f'<div style="margin-bottom:6px; font-size:{body_px}px;">{clean_markdown(line)}</div>'
        
        if in_table:
            html_body += '</table>'
        
        # 背景图
        cv_bg = ""
        bd_bg = ""
        if st.session_state.bg_cover_bytes:
            b64 = base64.b64encode(st.session_state.bg_cover_bytes).decode()
            cv_bg = f"background-image:url('data:image/png;base64,{b64}'); background-size:cover;"
        if st.session_state.bg_body_bytes:
            b64 = base64.b64encode(st.session_state.bg_body_bytes).decode()
            bd_bg = f"background-image:url('data:image/png;base64,{b64}'); background-size:cover;"
        
        preview_html = f"""
        <div style="background:#f1f5f9; padding:20px; border-radius:12px;">
            <div style="width:100%; max-width:794px; margin:0 auto; background:white; box-shadow:0 10px 40px rgba(0,0,0,0.2); border-radius:4px; overflow:hidden;">
                <div style="height:400px; display:flex; align-items:center; justify-content:center; {cv_bg} background-color:white;">
                    <h1 style="font-size:42px; color:#1e293b; text-align:center; font-weight:bold;">{cover_title}</h1>
                </div>
                <div style="padding:60px 50px; {bd_bg} background-color:white; font-family:'{font_family}', sans-serif; line-height:1.6;">
                    {html_body}
                </div>
            </div>
        </div>
        """
        st.components.v1.html(preview_html, height=700, scrolling=True)

# ═══════════════════════════════════════════
# 10. 右侧面板 - 图库与导出
# ═══════════════════════════════════════════
with col_right:
    st.markdown("### 📷 图库管理")
    
    # 图片上传
    with st.container(border=True):
        uploaded_imgs = st.file_uploader("上传图片", accept_multiple_files=True, type=['png','jpg','jpeg','webp'], label_visibility="collapsed")
        if uploaded_imgs:
            for f in uploaded_imgs:
                bytes_data = f.getvalue()
                st.session_state.gallery_images.append((bytes_data, f.name))
            log(f"上传 {len(uploaded_imgs)} 张图片", "success")
            st.rerun()
    
    # 图片网格
    if st.session_state.gallery_images:
        st.markdown(f"**共 {len(st.session_state.gallery_images)} 张图片**")
        
        # 全选/取消
        c1, c2 = st.columns(2)
        if c1.button("☐ 全选", use_container_width=True):
            st.session_state.selected_images = set(range(len(st.session_state.gallery_images)))
            st.rerun()
        if c2.button("🗑 清空", use_container_width=True):
            st.session_state.gallery_images = []
            st.session_state.selected_images = set()
            log("图库已清空", "warning")
            st.rerun()
        
        # 显示图片网格
        cols_per_row = 2
        for i in range(0, len(st.session_state.gallery_images), cols_per_row):
            cols = st.columns(cols_per_row)
            for j in range(cols_per_row):
                idx = i + j
                if idx < len(st.session_state.gallery_images):
                    img_bytes, img_name = st.session_state.gallery_images[idx]
                    is_selected = idx in st.session_state.selected_images
                    with cols[j]:
                        border_color = "#667eea" if is_selected else "transparent"
                        st.markdown(f'<div style="border:3px solid {border_color}; border-radius:8px; padding:4px;">', unsafe_allow_html=True)
                        st.image(img_bytes, use_container_width=True)
                        if st.checkbox("选择", key=f"img_sel_{idx}", value=is_selected):
                            st.session_state.selected_images.add(idx)
                        else:
                            st.session_state.selected_images.discard(idx)
                        st.markdown('</div>', unsafe_allow_html=True)
    else:
        st.info("暂无图片\n\n请上传或从文档提取")
    
    st.divider()
    
    # 背景图设置
    st.markdown("### 🖼️ 背景图")
    with st.container(border=True):
        cover_bg = st.file_uploader("封面背景", type=['png','jpg'], key="cover_bg_uploader")
        if cover_bg:
            st.session_state.bg_cover_bytes = cover_bg.getvalue()
            log("封面背景已设置", "success")
        
        body_bg = st.file_uploader("正文水印", type=['png','jpg'], key="body_bg_uploader")
        if body_bg:
            st.session_state.bg_body_bytes = body_bg.getvalue()
            log("正文水印已设置", "success")
        
        if st.button("🗑️ 清空背景", use_container_width=True):
            st.session_state.bg_cover_bytes = None
            st.session_state.bg_body_bytes = None
            log("背景已清空", "warning")
            st.rerun()
    
    st.divider()
    
    # 导出区
    st.markdown("### 📥 导出文档")
    with st.container(border=True):
        st.text_input("中文文件名", key="cover_cn", value=st.session_state.get('cover_cn', '产品规格书'))
        st.text_input("英文文件名", key="cover_en", value=st.session_state.get('cover_en', 'Product Specification'))
        
        if st.session_state.txt_cn:
            st.download_button(
                "📥 导出中文 Word",
                data=generate_word_document('cn'),
                file_name=f"{st.session_state.cover_cn}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
        
        if st.session_state.txt_en:
            st.download_button(
                "📥 导出英文 Word",
                data=generate_word_document('en'),
                file_name=f"{st.session_state.cover_en}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )

# ═══════════════════════════════════════════
# 11. 页脚
# ═══════════════════════════════════════════
st.divider()
st.markdown("""
<div style="text-align: center; color: #64748b; font-size: 12px; padding: 10px;">
    WPS 规格书助手 Web 全能版 V1.0 | 基于 Streamlit 构建 | 
    <a href="https://github.com/" target="_blank" style="color: #667eea;">GitHub 部署指南</a>
</div>
""", unsafe_allow_html=True)
'''

# 保存到文件
with open('/mnt/agents/output/app.py', 'w', encoding='utf-8') as f:
    f.write(code)

print("✅ 修复版代码已生成！")
print(f"📁 文件: /mnt/agents/output/app.py")
print(f"📊 长度: {len(code)} 字符")
